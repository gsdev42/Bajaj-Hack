from typing import List
from langchain.schema import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Temporary main function to avoid error
import requests
from io import BytesIO
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import mailparser
from dotenv import load_dotenv
load_dotenv()


def fetch_file_from_url(url):
    """Fetch file from URL and return file stream and content type."""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
    }
    response = requests.get(url, headers=headers, timeout=30)
    response.raise_for_status()
    
    content_type = response.headers.get("content-type", "").lower()
    return BytesIO(response.content), content_type


def detect_file_type(url, content_type):
    """Detect file type based on URL and content-type header."""
    url_lower = url.lower()
    content_lower = content_type.lower()

    if "pdf" in content_lower or url_lower.endswith(".pdf"):
        return "pdf"
    elif ("word" in content_lower or "document" in content_lower or 
          url_lower.endswith(".docx")):
        return "docx"
    elif "eml" in content_lower or url_lower.endswith(".eml"):
        return "eml"
    else:
        raise ValueError(f"Unsupported file type: {content_type}")


def parse_pdf(file_stream):
    """Extract text from a PDF file stream."""
    file_stream.seek(0)
    text = ""
    with fitz.open(stream=file_stream.read(), filetype="pdf") as doc:
        for page in doc:
            text += page.get_text() + "\n"
    return text.strip()


def parse_docx(file_stream):
    """Extract text from a DOCX file stream."""
    file_stream.seek(0)
    doc = DocxDocument(file_stream)
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    return "\n".join(paragraphs)


def parse_eml(file_stream):
    """Extract body text from an EML email file."""
    file_stream.seek(0)
    content = file_stream.read()
    mail = mailparser.parse_from_bytes(content)
    
    if mail.body:
        return mail.body
    elif hasattr(mail, 'text_plain') and mail.text_plain:
        return mail.text_plain[0] if isinstance(mail.text_plain, list) else mail.text_plain
    else:
        return ""


def parse_document_from_url(url: str) -> str:
    """Main entry: Fetch, detect type, and extract text from a document URL."""
    file_stream, content_type = fetch_file_from_url(url)
    file_type = detect_file_type(url, content_type)

    if file_type == "pdf":
        return parse_pdf(file_stream)
    elif file_type == "docx":
        return parse_docx(file_stream)
    elif file_type == "eml":
        return parse_eml(file_stream)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def parse_document(url: str) -> List[Document]:
    """Parse URL content into LangChain Documents"""
    text = parse_document_from_url(url)
    
    if not text:
        raise ValueError("No text extracted from document")
    
    # Create parent document
    doc = Document(
        page_content=text,
        metadata={"source": url, "total_chars": len(text)}
    )
    
    # Split into chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=400,
        length_function=len,
        separators=["\n\n", "\n", " ", ""]
    )
    
    return text_splitter.split_documents([doc])